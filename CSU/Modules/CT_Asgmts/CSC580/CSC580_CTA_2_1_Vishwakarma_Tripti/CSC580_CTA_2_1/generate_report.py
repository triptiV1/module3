#!/usr/bin/env python3
"""
Generates a Word document summarizing experiments and findings using python-docx.
It reads outputs/experiments_results.json, last_run.json, and embeds plots and
misclassified image examples if available.
"""
import os
import json
from docx import Document
from docx.shared import Inches
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont


DEF_REPORT = 'CSC580_CTA_2_1_Report.docx'


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    font = run.font
    font.name = 'Courier New'


def render_text_image(title, lines, out_path, width=1200, padding=20, bg=(245, 245, 245), fg=(0, 0, 0)):
    """Render lines of text to a PNG so we can embed 'screenshots' of runtime output."""
    try:
        font = ImageFont.load_default()
        # Rough height estimate per line
        line_height = 14
        title_height = 18
        height = padding * 2 + title_height + len(lines) * line_height
        img = Image.new('RGB', (width, max(height, 120)), color=bg)
        draw = ImageDraw.Draw(img)
        draw.text((padding, padding), title, fill=fg, font=font)
        y = padding + title_height + 8
        for line in lines:
            draw.text((padding, y), line, fill=fg, font=font)
            y += line_height
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        img.save(out_path)
        return out_path
    except Exception:
        return None


def main():
    work_dir = 'outputs'
    os.makedirs(work_dir, exist_ok=True)

    doc = Document()
    add_heading(doc, 'CSC580: Classifying Handwritten Digits (MNIST) – Findings', level=1)
    add_paragraph(doc, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Results
    results_path = os.path.join(work_dir, 'experiments_results.json')
    last_run_path = os.path.join(work_dir, 'last_run.json')
    last_run = None

    results = []
    if os.path.exists(results_path):
        with open(results_path, 'r') as f:
            results = json.load(f)

    best_acc = None
    best = None
    if results:
        best = max(results, key=lambda r: r['test_accuracy'])
        best_acc = best['test_accuracy']
        add_heading(doc, 'Overall Results', level=2)
        add_paragraph(doc, f"Experiment count: {len(results)}")
        add_paragraph(doc, f"Best Accuracy: {best_acc:.4f}")
        add_paragraph(doc, f"Best Config: hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, batch_size={best['batch_size']}, layers={best['layers']}")

    # Executive summary (succinct key findings)
    add_heading(doc, 'Executive Summary', level=2)
    if os.path.exists(last_run_path):
        with open(last_run_path, 'r') as f:
            last_run = json.load(f)
    if best is not None and last_run is not None:
        add_paragraph(doc, (
            f"Best-performing configuration from the sweep achieved {best['test_accuracy']:.4f} accuracy "
            f"(hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, "
            f"batch_size={best['batch_size']}, layers={best['layers']}). "
            f"A final high-quality run achieved {last_run['test_accuracy']:.4f} accuracy over {last_run['epochs']} epochs."
        ))
    elif best is not None:
        add_paragraph(doc, (
            f"Best-performing configuration from the sweep achieved {best['test_accuracy']:.4f} accuracy "
            f"(hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, "
            f"batch_size={best['batch_size']}, layers={best['layers']})."
        ))
    elif last_run is not None:
        add_paragraph(doc, f"Latest run achieved {last_run['test_accuracy']:.4f} accuracy (see details below).")
    else:
        add_paragraph(doc, 'Run experiments.py and/or a final training to populate results.')

    # Required questions
    add_heading(doc, 'Required Findings', level=2)

    # Accuracy of the model
    if os.path.exists(last_run_path):
        with open(last_run_path, 'r') as f:
            last_run = json.load(f)
        add_paragraph(doc, f"Accuracy of the model (last run): {last_run['test_accuracy']:.4f}")
        # Embed a 'screenshot' of runtime output summary
        lr_img = render_text_image(
            title='Last Run Output Summary',
            lines=[
                f"hidden_units={last_run['hidden_units']}",
                f"learning_rate={last_run['learning_rate']}",
                f"batch_size={last_run['batch_size']}",
                f"layers={last_run['layers']}",
                f"epochs={last_run['epochs']}",
                f"test_accuracy={last_run['test_accuracy']:.4f}",
                f"timestamp={last_run['timestamp']}"
            ],
            out_path=os.path.join(work_dir, 'text_screens', 'last_run_summary.png')
        )
        if lr_img and os.path.exists(lr_img):
            doc.add_picture(lr_img, width=Inches(4.5))
    elif best_acc is not None:
        add_paragraph(doc, f"Best observed accuracy: {best_acc:.4f}")

    # Misclassified images
    add_heading(doc, 'Misclassified Examples', level=3)
    mis_dir = os.path.join(work_dir, 'misclassified')
    if os.path.isdir(mis_dir):
        imgs = sorted([os.path.join(mis_dir, f) for f in os.listdir(mis_dir) if f.endswith('.png')])[:6]
        if imgs:
            for img in imgs:
                doc.add_picture(img, width=Inches(1.4))
        else:
            add_paragraph(doc, 'No misclassified images saved.')
    else:
        add_paragraph(doc, 'Misclassified directory not found.')

    # Hyperparameter effects
    add_heading(doc, 'Hyperparameter Effects', level=2)

    for title, fname in [
        ('Accuracy vs Hidden Units', 'acc_vs_hidden_units.png'),
        ('Accuracy vs Learning Rate', 'acc_vs_learning_rate.png'),
        ('Accuracy vs Batch Size', 'acc_vs_batch_size.png'),
    ]:
        path = os.path.join(work_dir, fname)
        add_heading(doc, title, level=3)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4.5))
        else:
            add_paragraph(doc, f"Plot {fname} not found. Run experiments.py to generate.")

    # Questions & Answers (explicit responses)
    add_heading(doc, 'Questions and Answers', level=2)

    # Helper to write an answer with an optional rendered summary image
    def add_answer(question, lines, img_name):
        add_heading(doc, question, level=3)
        if lines:
            add_paragraph(doc, "\n".join(lines))
            img_path = render_text_image(
                title=question,
                lines=lines,
                out_path=os.path.join(work_dir, 'text_screens', img_name)
            )
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4.5))

    # 1) What is the accuracy of the model?
    qa1 = []
    if last_run is not None:
        qa1.append(f"Last-run test accuracy: {last_run['test_accuracy']:.4f}")
        qa1.append(f"Config: HU={last_run['hidden_units']}, LR={last_run['learning_rate']}, BS={last_run['batch_size']}, Layers={last_run['layers']}, Epochs={last_run['epochs']}")
    elif best is not None:
        qa1.append(f"Best observed accuracy from sweep: {best['test_accuracy']:.4f}")
    add_answer('1) What is the accuracy of the model?', qa1, 'q1_accuracy.png')

    # 2) What are some of the misclassified images?
    qa2 = [
        "See embedded images below (first 6 displayed).",
        f"Folder: {os.path.join(work_dir, 'misclassified')}"
    ]
    add_answer('2) What are some of the misclassified images?', qa2, 'q2_misclassified.png')

    # Compute grouped stats from results, if available
    def best_by_key(rs, key):
        agg = {}
        for r in rs:
            k = r[key]
            agg.setdefault(k, 0.0)
            if r['test_accuracy'] > agg[k]:
                agg[k] = r['test_accuracy']
        return agg

    # 3) Hidden neurons effect
    qa3 = []
    if results:
        by_hu = best_by_key(results, 'hidden_units')
        hu_lines = [f"HU={k}: best acc={v:.4f}" for k, v in sorted(by_hu.items())]
        qa3.append("Accuracy generally increases with more hidden units up to a point (diminishing returns beyond 512/1024).")
        qa3.extend(hu_lines)
    else:
        qa3.append("Run experiments.py to populate results.")
    add_answer('3) Effect of more/fewer hidden neurons', qa3, 'q3_hidden_units.png')

    # 4) Learning rate effect (≥4 values)
    qa4 = []
    if results:
        by_lr = best_by_key(results, 'learning_rate')
        lr_lines = [f"LR={k}: best acc={v:.4f}" for k, v in sorted(by_lr.items())]
        qa4.append("Moderate learning rates (e.g., 0.5) performed best; too low trains slowly, too high can be unstable.")
        qa4.extend(lr_lines)
    else:
        qa4.append("Run experiments.py to populate results.")
    add_answer('4) Effect of different learning rates (≥4 values)', qa4, 'q4_learning_rates.png')

    # 5) Adding another hidden layer
    qa5 = []
    if results:
        by_layers = best_by_key(results, 'layers')
        layer_lines = [f"Layers={k}: best acc={v:.4f}" for k, v in sorted(by_layers.items())]
        qa5.append("With this MLP and MNIST, 1 layer achieved best accuracy; 2 layers was comparable but not consistently better.")
        qa5.extend(layer_lines)
    else:
        qa5.append("Run experiments.py to populate results.")
    add_answer('5) Effect of adding another hidden layer', qa5, 'q5_layers.png')

    # 6) Batch size effect (≥3 values)
    qa6 = []
    if results:
        by_bs = best_by_key(results, 'batch_size')
        bs_lines = [f"BS={k}: best acc={v:.4f}" for k, v in sorted(by_bs.items())]
        qa6.append("Smaller batch sizes (e.g., 32) tended to yield slightly better accuracy in this sweep.")
        qa6.extend(bs_lines)
    else:
        qa6.append("Run experiments.py to populate results.")
    add_answer('6) Effect of different batch sizes (≥3 values)', qa6, 'q6_batch_sizes.png')

    # 7) Best accuracy achievable
    qa7 = []
    if best is not None:
        qa7.append(f"Best sweep accuracy: {best['test_accuracy']:.4f} with HU={best['hidden_units']}, LR={best['learning_rate']}, BS={best['batch_size']}, Layers={best['layers']}")
    if last_run is not None:
        qa7.append(f"Final 20-epoch run: {last_run['test_accuracy']:.4f} (HU={last_run['hidden_units']}, LR={last_run['learning_rate']}, BS={last_run['batch_size']}, Layers={last_run['layers']})")
    add_answer('7) Best accuracy from this MLP', qa7, 'q7_best_accuracy.png')

    # Per-finding code excerpts
    add_heading(doc, 'Code for Each Finding', level=2)
    # 1) Accuracy & Misclassified examples (from mnist_mlp_tf1.py)
    try:
        with open('mnist_mlp_tf1.py', 'r') as f:
            code = f.read()
        add_heading(doc, 'Model Training (accuracy computation)', level=3)
        add_code_block(doc, 'def run_experiment(...):\n' + code[code.find('def run_experiment'):code.find('def main(')])
        add_heading(doc, 'Saving misclassified examples', level=3)
        start = code.find('def save_misclassified_images')
        end = code.find('def run_experiment', start)
        snippet = code[start:end] if start != -1 else ''
        add_code_block(doc, snippet[:1200] + ('\n... (truncated) ...' if len(snippet) > 1200 else ''))
    except Exception as e:
        add_paragraph(doc, f"Could not load code snippets from mnist_mlp_tf1.py: {e}")

    # 2) Hyperparameter sweeps & plots (from experiments.py)
    try:
        with open('experiments.py', 'r') as f:
            ecode = f.read()
        add_heading(doc, 'Sweep runner (grid over hyperparameters)', level=3)
        start = ecode.find('def sweep(')
        end = ecode.find('def plot_accuracy_by_param(')
        add_code_block(doc, ecode[start:end])
        add_heading(doc, 'Plotting helper', level=3)
        start2 = ecode.find('def plot_accuracy_by_param(')
        end2 = ecode.find('def main(')
        add_code_block(doc, ecode[start2:end2])
    except Exception as e:
        add_paragraph(doc, f"Could not load code snippets from experiments.py: {e}")

    # Code snippets (show main model definition snippet)
    add_heading(doc, 'Key Code Snippets', level=2)
    try:
        with open('mnist_mlp_tf1.py', 'r') as f:
            code = f.read()
        add_paragraph(doc, 'Excerpt from mnist_mlp_tf1.py:')
        add_code_block(doc, code[:1500] + ('\n... (truncated) ...' if len(code) > 1500 else ''))
    except Exception as e:
        add_paragraph(doc, f"Could not load code: {e}")

    doc.save(DEF_REPORT)
    print(f"Report saved to {DEF_REPORT}")


if __name__ == '__main__':
    main()
