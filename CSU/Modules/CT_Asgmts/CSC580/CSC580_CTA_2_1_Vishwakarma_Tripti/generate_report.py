#!/usr/bin/env python3
"""
Generates a Word document summarizing experiments and findings using python-docx.
It reads outputs/experiments_results.json, last_run.json, and embeds plots and
misclassified image examples if available.
"""
import os
import json
from docx import Document
from docx.shared import Inches
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont


DEF_REPORT = 'CSC580_CTA_2_1_Report.docx'


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    font = run.font
    font.name = 'Courier New'


def render_text_image(title, lines, out_path, width=1200, padding=20, bg=(245, 245, 245), fg=(0, 0, 0)):
    """Render lines of text to a PNG so we can embed 'screenshots' of runtime output."""
    try:
        font = ImageFont.load_default()
        # Rough height estimate per line
        line_height = 14
        title_height = 18
        height = padding * 2 + title_height + len(lines) * line_height
        img = Image.new('RGB', (width, max(height, 120)), color=bg)
        draw = ImageDraw.Draw(img)
        draw.text((padding, padding), title, fill=fg, font=font)
        y = padding + title_height + 8
        for line in lines:
            draw.text((padding, y), line, fill=fg, font=font)
            y += line_height
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        img.save(out_path)
        return out_path
    except Exception:
        return None


def main():
    work_dir = 'outputs'
    os.makedirs(work_dir, exist_ok=True)

    doc = Document()
    add_heading(doc, 'CSC580: Classifying Handwritten Digits (MNIST) – Findings', level=1)
    add_paragraph(doc, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Project Overview (for standalone submission without README)
    add_heading(doc, 'Project Overview', level=2)
    add_paragraph(doc, 'This project trains a multi-layer perceptron (MLP) on the MNIST dataset using TensorFlow in TF1-style via tf.compat.v1. It includes scripts to run hyperparameter experiments, save misclassified images, create plots, and generate this Word report.')
    add_heading(doc, 'Files Included', level=3)
    add_paragraph(doc, '- mnist_mlp_tf1.py – Core training script; can be run directly or imported for experiments.')
    add_paragraph(doc, '- experiments.py – Runs sweeps over hidden units, learning rates, batch sizes, and layers; saves plots and JSON results under outputs/.')
    add_paragraph(doc, '- generate_report.py – Builds this Word report with findings, plots, Q&A, and placeholders for screenshots.')
    add_paragraph(doc, '- requirements.txt – Python dependencies for Apple Silicon (tensorflow-macos + tensorflow-metal) and libraries (numpy, matplotlib, python-docx, scikit-learn, pandas).')
    add_paragraph(doc, '- outputs/ – Directory created by experiments and training containing plots, misclassified samples, and JSON logs.')
    add_heading(doc, 'Quick Start (summarized)', level=3)
    add_paragraph(doc, '1) Create/activate a virtual environment (Python 3.11 on Apple Silicon).')
    add_paragraph(doc, '2) pip install -r requirements.txt')
    add_paragraph(doc, '3) Single run (baseline): python mnist_mlp_tf1.py --hidden_units 512 --learning_rate 0.5 --batch_size 32 --epochs 20 --layers 1')
    add_paragraph(doc, '4) Run hyperparameter sweep: python experiments.py (plots and results saved to outputs/)')
    add_paragraph(doc, '5) Generate this report: python generate_report.py')
    add_heading(doc, 'Notes', level=3)
    add_paragraph(doc, '- Implementation follows TF1 placeholder/session pattern using tf.compat.v1.disable_eager_execution().')
    add_paragraph(doc, '- Default sweep epochs are reduced for speed; final high-quality run uses 20 epochs.')
    add_paragraph(doc, '- On Apple Silicon, use tensorflow-macos 2.15.0 + tensorflow-metal 1.1.0 for compatibility.')

    # Submission Details (fill in your info before submitting)
    add_heading(doc, 'Submission Details', level=2)
    add_paragraph(doc, 'Student: <Your Name Here>')
    add_paragraph(doc, 'Course: CSC580')
    add_paragraph(doc, 'Assignment: CTA 2.1 – MNIST MLP')
    add_paragraph(doc, 'Date: <Submission Date>')

    # Results
    results_path = os.path.join(work_dir, 'experiments_results.json')
    last_run_path = os.path.join(work_dir, 'last_run.json')
    last_run = None

    results = []
    if os.path.exists(results_path):
        with open(results_path, 'r') as f:
            results = json.load(f)

    best_acc = None
    best = None
    if results:
        best = max(results, key=lambda r: r['test_accuracy'])
        best_acc = best['test_accuracy']
        add_heading(doc, 'Overall Results', level=2)
        add_paragraph(doc, f"Experiment count: {len(results)}")
        add_paragraph(doc, f"Best Accuracy: {best_acc:.4f}")
        add_paragraph(doc, f"Best Config: hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, batch_size={best['batch_size']}, layers={best['layers']}")

    # Executive summary (succinct key findings)
    add_heading(doc, 'Executive Summary', level=2)
    if os.path.exists(last_run_path):
        with open(last_run_path, 'r') as f:
            last_run = json.load(f)
    if best is not None and last_run is not None:
        add_paragraph(doc, (
            f"Best-performing configuration from the sweep achieved {best['test_accuracy']:.4f} accuracy "
            f"(hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, "
            f"batch_size={best['batch_size']}, layers={best['layers']}). "
            f"A final high-quality run achieved {last_run['test_accuracy']:.4f} accuracy over {last_run['epochs']} epochs."
        ))
    elif best is not None:
        add_paragraph(doc, (
            f"Best-performing configuration from the sweep achieved {best['test_accuracy']:.4f} accuracy "
            f"(hidden_units={best['hidden_units']}, learning_rate={best['learning_rate']}, "
            f"batch_size={best['batch_size']}, layers={best['layers']})."
        ))
    elif last_run is not None:
        add_paragraph(doc, f"Latest run achieved {last_run['test_accuracy']:.4f} accuracy (see details below).")
    else:
        add_paragraph(doc, 'Run experiments.py and/or a final training to populate results.')

    # Required questions
    add_heading(doc, 'Required Findings', level=2)

    # Accuracy of the model
    if os.path.exists(last_run_path):
        with open(last_run_path, 'r') as f:
            last_run = json.load(f)
        add_paragraph(doc, f"Accuracy of the model (last run): {last_run['test_accuracy']:.4f}")
        add_paragraph(doc, 'Insert screenshot here: Terminal output of the final run showing Epoch logs and Final Test Accuracy.')
        # User will paste actual terminal screenshots; no auto-generated images.
    elif best_acc is not None:
        add_paragraph(doc, f"Best observed accuracy: {best_acc:.4f}")

    # Misclassified images (placeholder only; user will paste screenshots)
    add_heading(doc, 'Misclassified Examples', level=3)
    add_paragraph(doc, 'Insert screenshot here: A grid or multiple thumbnails from outputs/misclassified/*.png (e.g., 6 images).')

    # Hyperparameter effects
    add_heading(doc, 'Hyperparameter Effects', level=2)

    for title, fname in [
        ('Accuracy vs Hidden Units', 'acc_vs_hidden_units.png'),
        ('Accuracy vs Learning Rate', 'acc_vs_learning_rate.png'),
        ('Accuracy vs Batch Size', 'acc_vs_batch_size.png'),
    ]:
        path = os.path.join(work_dir, fname)
        add_heading(doc, title, level=3)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4.5))
        else:
            add_paragraph(doc, f"Plot {fname} not found. Run experiments.py to generate.")

    # Questions & Answers (explicit responses)
    add_heading(doc, 'Questions and Answers', level=2)

    # Helper to write an answer (no auto screenshots; user will paste real screenshots)
    def add_answer(question, lines, img_name):
        add_heading(doc, question, level=3)
        if lines:
            add_paragraph(doc, "\n".join(lines))

    # 1) What is the accuracy of the model?
    qa1 = []
    if last_run is not None:
        qa1.append(f"Last-run test accuracy: {last_run['test_accuracy']:.4f}")
        qa1.append(f"Config: HU={last_run['hidden_units']}, LR={last_run['learning_rate']}, BS={last_run['batch_size']}, Layers={last_run['layers']}, Epochs={last_run['epochs']}")
    elif best is not None:
        qa1.append(f"Best observed accuracy from sweep: {best['test_accuracy']:.4f}")
    add_answer('1) What is the accuracy of the model?', qa1, 'q1_accuracy.png')
    add_paragraph(doc, 'Insert screenshot here: Terminal output of the final 20-epoch run showing Final Test Accuracy.')

    # 2) What are some of the misclassified images?
    qa2 = [
        "See embedded images below (first 6 displayed).",
        f"Folder: {os.path.join(work_dir, 'misclassified')}"
    ]
    add_answer('2) What are some of the misclassified images?', qa2, 'q2_misclassified.png')
    add_paragraph(doc, 'Insert screenshot here: A grid or multiple thumbnails from outputs/misclassified/*.png (e.g., 6 images).')

    # Compute grouped stats from results, if available
    def best_by_key(rs, key):
        agg = {}
        for r in rs:
            k = r[key]
            agg.setdefault(k, 0.0)
            if r['test_accuracy'] > agg[k]:
                agg[k] = r['test_accuracy']
        return agg

    # 3) Hidden neurons effect
    qa3 = []
    if results:
        by_hu = best_by_key(results, 'hidden_units')
        hu_lines = [f"HU={k}: best acc={v:.4f}" for k, v in sorted(by_hu.items())]
        qa3.append("Accuracy generally increases with more hidden units up to a point (diminishing returns beyond 512/1024).")
        qa3.extend(hu_lines)
    else:
        qa3.append("Run experiments.py to populate results.")
    add_answer('3) Effect of more/fewer hidden neurons', qa3, 'q3_hidden_units.png')
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_hidden_units.png plot.')

    # 4) Learning rate effect (≥4 values)
    qa4 = []
    if results:
        by_lr = best_by_key(results, 'learning_rate')
        lr_lines = [f"LR={k}: best acc={v:.4f}" for k, v in sorted(by_lr.items())]
        qa4.append("Moderate learning rates (e.g., 0.5) performed best; too low trains slowly, too high can be unstable.")
        qa4.extend(lr_lines)
    else:
        qa4.append("Run experiments.py to populate results.")
    add_answer('4) Effect of different learning rates (≥4 values)', qa4, 'q4_learning_rates.png')
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_learning_rate.png plot.')

    # 5) Adding another hidden layer
    qa5 = []
    if results:
        by_layers = best_by_key(results, 'layers')
        layer_lines = [f"Layers={k}: best acc={v:.4f}" for k, v in sorted(by_layers.items())]
        qa5.append("With this MLP and MNIST, 1 layer achieved best accuracy; 2 layers was comparable but not consistently better.")
        qa5.extend(layer_lines)
    else:
        qa5.append("Run experiments.py to populate results.")
    add_answer('5) Effect of adding another hidden layer', qa5, 'q5_layers.png')
    add_paragraph(doc, 'Insert screenshot here: Small snippet showing comparison where layers=1 vs layers=2 (terminal tail or best_config.json excerpt).')

    # 6) Batch size effect (≥3 values)
    qa6 = []
    if results:
        by_bs = best_by_key(results, 'batch_size')
        bs_lines = [f"BS={k}: best acc={v:.4f}" for k, v in sorted(by_bs.items())]
        qa6.append("Smaller batch sizes (e.g., 32) tended to yield slightly better accuracy in this sweep.")
        qa6.extend(bs_lines)
    else:
        qa6.append("Run experiments.py to populate results.")
    add_answer('6) Effect of different batch sizes (≥3 values)', qa6, 'q6_batch_sizes.png')
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_batch_size.png plot.')

    # 7) Best accuracy achievable
    qa7 = []
    if best is not None:
        qa7.append(f"Best sweep accuracy: {best['test_accuracy']:.4f} with HU={best['hidden_units']}, LR={best['learning_rate']}, BS={best['batch_size']}, Layers={best['layers']}")
    if last_run is not None:
        qa7.append(f"Final 20-epoch run: {last_run['test_accuracy']:.4f} (HU={last_run['hidden_units']}, LR={last_run['learning_rate']}, BS={last_run['batch_size']}, Layers={last_run['layers']})")
    add_answer('7) Best accuracy from this MLP', qa7, 'q7_best_accuracy.png')
    add_paragraph(doc, 'Insert screenshot here: Terminal line showing Best configuration (from experiments.py) and Final Test Accuracy from the 20-epoch run.')

    # Per-finding code excerpts
    add_heading(doc, 'Code for Each Finding', level=2)
    # 1) Accuracy & Misclassified examples (from mnist_mlp_tf1.py)
    try:
        with open('mnist_mlp_tf1.py', 'r') as f:
            code = f.read()
        add_heading(doc, 'Model Training (accuracy computation)', level=3)
        add_code_block(doc, 'def run_experiment(...):\n' + code[code.find('def run_experiment'):code.find('def main(')])
        add_heading(doc, 'Saving misclassified examples', level=3)
        start = code.find('def save_misclassified_images')
        end = code.find('def run_experiment', start)
        snippet = code[start:end] if start != -1 else ''
        add_code_block(doc, snippet[:1200] + ('\n... (truncated) ...' if len(snippet) > 1200 else ''))
    except Exception as e:
        add_paragraph(doc, f"Could not load code snippets from mnist_mlp_tf1.py: {e}")

    # 2) Hyperparameter sweeps & plots (from experiments.py)
    try:
        with open('experiments.py', 'r') as f:
            ecode = f.read()
        add_heading(doc, 'Sweep runner (grid over hyperparameters)', level=3)
        start = ecode.find('def sweep(')
        end = ecode.find('def plot_accuracy_by_param(')
        add_code_block(doc, ecode[start:end])
        add_heading(doc, 'Plotting helper', level=3)
        start2 = ecode.find('def plot_accuracy_by_param(')
        end2 = ecode.find('def main(')
        add_code_block(doc, ecode[start2:end2])
    except Exception as e:
        add_paragraph(doc, f"Could not load code snippets from experiments.py: {e}")

    # Code snippets (show main model definition snippet)
    add_heading(doc, 'Key Code Snippets', level=2)
    try:
        with open('mnist_mlp_tf1.py', 'r') as f:
            code = f.read()
        add_paragraph(doc, 'Excerpt from mnist_mlp_tf1.py:')
        add_code_block(doc, code[:1500] + ('\n... (truncated) ...' if len(code) > 1500 else ''))
    except Exception as e:
        add_paragraph(doc, f"Could not load code: {e}")

    # Appendix with screenshot placeholders
    add_heading(doc, 'Appendix: Screenshot Placeholders', level=1)
    add_paragraph(doc, 'Paste your actual screenshots in the spaces below. Keep each item on its own line for clarity:')
    add_heading(doc, 'A1) Final Run Accuracy (Terminal)', level=2)
    add_paragraph(doc, 'Insert screenshot here: Terminal output showing Final Test Accuracy from the 20-epoch run.')
    add_heading(doc, 'A2) Best Configuration (Terminal/JSON)', level=2)
    add_paragraph(doc, 'Insert screenshot here: Terminal line "Best configuration: ..." or an excerpt of outputs/best_config.json.')
    add_heading(doc, 'A3) Hidden Units Plot', level=2)
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_hidden_units.png')
    add_heading(doc, 'A4) Learning Rate Plot', level=2)
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_learning_rate.png')
    add_heading(doc, 'A5) Batch Size Plot', level=2)
    add_paragraph(doc, 'Insert screenshot here: outputs/acc_vs_batch_size.png')
    add_heading(doc, 'A6) Layers Comparison', level=2)
    add_paragraph(doc, 'Insert screenshot here: Comparison snippet showing layers=1 vs layers=2 results.')
    add_heading(doc, 'A7) Misclassified Examples Grid', level=2)
    add_paragraph(doc, 'Insert screenshot here: A grid of several images from outputs/misclassified/.')

    doc.save(DEF_REPORT)
    print(f"Report saved to {DEF_REPORT}")


if __name__ == '__main__':
    main()
